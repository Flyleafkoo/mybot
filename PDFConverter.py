import logging
import os
from sklearn.cluster import KMeans
import pdfplumber
from pdf2image import convert_from_path
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np


def find_column_index(left, column_boundaries):
    for i, boundary in enumerate(column_boundaries):
        if left < boundary:
            return i
    return len(column_boundaries) - 1


class PDFConverter:
    def __init__(self, language='rus'):
        self.language = language

    def pdf_to_word(self, pdf_path, word_path):
        try:
            logging.info(f"Начало конвертации PDF в изображения: {pdf_path}")
            images = convert_from_path(pdf_path)
            doc = Document()

            for i, image in enumerate(images):
                logging.info(f"Обработка страницы {i + 1}")
                data = pytesseract.image_to_data(image, lang=self.language, output_type=pytesseract.Output.DICT)

                previous_bottom = 0
                for j in range(len(data['text'])):
                    if int(data['conf'][j]) > 60:  # Уверенность OCR
                        text = data['text'][j].strip()
                        if text:
                            top = data['top'][j]

                            # Добавляем новый параграф, если текст находится ниже предыдущего
                            if top > previous_bottom + 10:
                                paragraph = doc.add_paragraph()
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                run = paragraph.add_run(text)
                                run.font.size = Pt(10)
                            else:
                                # Добавляем текст в текущую строку
                                paragraph = doc.paragraphs[-1]
                                run = paragraph.add_run(" " + text)
                                run.font.size = Pt(10)

                            previous_bottom = top + data['height'][j]

                # Добавляем пустую строку между страницами
                doc.add_paragraph()

                logging.info(f"Текст извлечён с страницы {i + 1}")

            directory = os.path.dirname(word_path)
            if directory and not os.path.exists(directory):
                os.makedirs(directory)
                logging.info(f"Директория {directory} была создана.")

            doc.save(word_path)
            logging.info(f"Файл Word успешно сохранён: {word_path}")

        except Exception as e:
            logging.error(f"Ошибка при конвертации PDF в Word: {e}")

    def define_column_boundaries(self, positions, n_clusters=5):
        positions_array = np.array(positions).reshape(-1, 1)
        kmeans = KMeans(n_clusters=min(n_clusters, len(positions)), random_state=0)
        kmeans.fit(positions_array)
        return sorted(kmeans.cluster_centers_.flatten())

    def pdf_to_excel(self, pdf_path, excel_path):
        try:
            wb = Workbook()
            ws = wb.active

            thin_border = Border(left=Side(style='thin'),
                                 right=Side(style='thin'),
                                 top=Side(style='thin'),
                                 bottom=Side(style='thin'))

            row_offset = 0

            with pdfplumber.open(pdf_path) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    logging.info(f"Обработка страницы {page_number}")

                    # Извлечение таблиц с помощью pdfplumber
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                ws.append(row)
                                for col_index, _ in enumerate(row, start=1):
                                    cell = ws.cell(row=row_offset + 1, column=col_index)
                                    cell.font = Font(size=10, bold=True if row_offset == 0 else False)
                                    cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                                    cell.border = thin_border
                                row_offset += 1
                            row_offset += 1  # Отступ между таблицами

                    # OCR для извлечения текста со сканированных страниц
                    images = convert_from_path(pdf_path, first_page=page_number, last_page=page_number)
                    for image in images:
                        data = pytesseract.image_to_data(image, lang=self.language, output_type=pytesseract.Output.DICT)

                        # Определение границ колонок через K-Means
                        left_positions = [data['left'][i] for i in range(len(data['text'])) if
                                          int(data['conf'][i]) > 60]
                        if left_positions:
                            column_boundaries = self.define_column_boundaries(left_positions)

                            current_row = [''] * len(column_boundaries)
                            previous_bottom = 0

                            for i in range(len(data['text'])):
                                if int(data['conf'][i]) > 60:  # Уверенность OCR
                                    text = data['text'][i].strip()
                                    if text:
                                        left = data['left'][i]
                                        top = data['top'][i]

                                        # Определение колонки по позиции
                                        column_index = find_column_index(left, column_boundaries)

                                        # Проверка на начало новой строки
                                        if top > previous_bottom + 10:
                                            if any(current_row):
                                                ws.append(current_row)
                                                for col_index in range(len(current_row)):
                                                    cell = ws.cell(row=row_offset + 1, column=col_index + 1)
                                                    cell.font = Font(size=10)
                                                    cell.alignment = Alignment(horizontal='left', vertical='top',
                                                                               wrap_text=True)
                                                    cell.border = thin_border
                                                row_offset += 1
                                            current_row = [''] * len(column_boundaries)

                                        current_row[column_index] += (' ' + text).strip()
                                        previous_bottom = top + data['height'][i]

                            # Добавление последней строки
                            if any(current_row):
                                ws.append(current_row)
                                for col_index in range(len(current_row)):
                                    cell = ws.cell(row=row_offset + 1, column=col_index + 1)
                                    cell.font = Font(size=10)
                                    cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                                    cell.border = thin_border
                                row_offset += 1

            # Настройка ширины колонок
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter  # Получение буквы колонки
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column].width = adjusted_width

            wb.save(excel_path)
            logging.info(f"Файл Excel успешно сохранён: {excel_path}")

        except Exception as e:
            logging.error(f"Ошибка при конвертации PDF в Excel: {e}")
